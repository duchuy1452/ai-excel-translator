from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from typing import Callable, List, Tuple
from io import BytesIO
import math
import streamlit as st

# Import necessary constants
from constants import CHARS_PER_TOKEN, MAX_TOKENS_PER_REQUEST
from translators import GeminiTranslator
from utils import logger


class DocxProcessor:
    """
    Processes Word (.docx) files, extracts text from paragraphs and tables,
    translates it using a translator (with token-based batching), and creates
    a new Word file with the translated text.
    """

    def __init__(self, translator: GeminiTranslator):
        self.translator = translator

    @staticmethod
    def estimate_tokens(text: str) -> int:
        """Estimates the number of tokens in a text."""
        # Use math.ceil to slightly overestimate, ensuring batches don't exceed limit
        return (
            math.ceil(len(text) / CHARS_PER_TOKEN) if CHARS_PER_TOKEN > 0 else len(text)
        )

    def create_batches(
        self, items_with_indices: List[Tuple[int, str]]
    ) -> List[List[Tuple[int, str]]]:
        """
        Creates batches of (original_index, text) tuples based on the token limit.

        Args:
            items_with_indices: A list of tuples, where each tuple is (original_index, text_to_translate).

        Returns:
            A list of batches, where each batch is a list of (original_index, text) tuples.
        """
        batches: List[List[Tuple[int, str]]] = []
        current_batch: List[Tuple[int, str]] = []
        current_token_count = 0

        for index, text in items_with_indices:
            token_count = self.estimate_tokens(text)
            if (
                current_token_count + token_count > MAX_TOKENS_PER_REQUEST
                and current_batch
            ):
                batches.append(current_batch)
                current_batch = []
                current_token_count = 0

            current_batch.append((index, text))
            current_token_count += token_count

        if current_batch:
            batches.append(current_batch)

        return batches

    def translate_paragraphs(
        self, document: Document, target_lang: str, progress_callback: Callable = None
    ):
        paragraphs: List[Paragraph] = [p for p in document.paragraphs]
        paragraph_texts_with_indices: List[Tuple[int, str]] = []
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if text and not text.isnumeric():
                paragraph_texts_with_indices.append((i, text))

        paragraph_batches = self.create_batches(paragraph_texts_with_indices)
        total_paragraph_batches = len(paragraph_batches)
        for i, batch in enumerate(paragraph_batches):
            texts_to_translate = [text for index, text in batch]
            translated_batch = self.translator.translate_batch(
                texts_to_translate, target_lang=target_lang
            )
            for index, translated_text in zip(
                [index for index, text in batch], translated_batch
            ):
                paragraphs[index].text = translated_text
            if progress_callback:
                progress_callback(i + 1, total_paragraph_batches)

    def translate_tables(
        self, document: Document, target_lang: str, progress_callback: Callable = None
    ):
        total_tables = len(document.tables)
        for table_index, table in enumerate(document.tables):
            table_texts_with_indices: List[Tuple[int, str]] = []
            cells: List[_Cell] = []
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text and not text.isnumeric():
                        table_texts_with_indices.append((len(cells), text))
                        cells.append(cell)

            table_batches = self.create_batches(table_texts_with_indices)
            for batch in table_batches:
                texts_to_translate = [text for index, text in batch]
                translated_batch = self.translator.translate_batch(
                    texts_to_translate, target_lang=target_lang
                )
                for index, translated_text in zip(
                    [index for index, text in batch], translated_batch
                ):
                    cells[index].text = translated_text
            if progress_callback:
                progress_callback(table_index + 1, total_tables)

    def process_docx(
        self, file_data: BytesIO, target_lang: str, progress_callback: Callable = None
    ) -> BytesIO:
        """
        Processes a Word file, translates its content using token-based batching,
        and returns the translated file as a BytesIO stream.
        """
        document = Document(file_data)

        total_items = len(document.paragraphs) + len(document.tables)

        logger.info("Translating paragraphs")
        st.write("1. Translating paragraphs")
        self.translate_paragraphs(document, target_lang, progress_callback)
        logger.info("Translating tables")
        st.write("2. Translating tables")
        self.translate_tables(document, target_lang, progress_callback)

        # Save the translated document to BytesIO
        output_file = BytesIO()
        document.save(output_file)
        output_file.seek(0)  # Reset the buffer's position to the beginning
        return output_file
